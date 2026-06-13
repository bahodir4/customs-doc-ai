"""Tests for the customs-law ingestion service."""
from __future__ import annotations

from pathlib import Path

import pytest
import pytest_asyncio

from rag import LexIngestionService, UnsupportedSourceError


# ── Unit-style tests ─────────────────────────────────────────────────


class TestSourceDetection:
    @pytest.mark.parametrize("source, expected", [
        ("http://example.com", True),
        ("https://lex.uz/docs/1", True),
        ("HTTP://EXAMPLE.COM", True),
        ("file.docx", False),
        ("/abs/path.md", False),
        ("not-a-url", False),
    ])
    def test_is_url(self, source: str, expected: bool) -> None:
        assert LexIngestionService._is_url(source) is expected


@pytest.mark.asyncio
async def test_unsupported_extension_raises(tmp_path: Path) -> None:
    bad = tmp_path / "notes.xyz"
    bad.write_text("hello")

    # We can construct the service without a real vector store because
    # the dispatch check happens before any vector calls.
    service = LexIngestionService(vector_store=None)  # type: ignore[arg-type]
    with pytest.raises(UnsupportedSourceError, match="Unsupported"):
        await service.ingest(bad)


@pytest.mark.asyncio
async def test_missing_file_raises(tmp_path: Path) -> None:
    service = LexIngestionService(vector_store=None)  # type: ignore[arg-type]
    with pytest.raises(FileNotFoundError):
        await service.ingest(tmp_path / "ghost.md")


# ── Integration tests (require live Qdrant + Ollama for embeddings) ──


@pytest_asyncio.fixture
async def ingest_service(vector_store):
    return LexIngestionService(vector_store=vector_store)


@pytest.fixture
def sample_md(tmp_path: Path) -> Path:
    """A small markdown file with a clear hierarchy."""
    md_path = tmp_path / "customs.md"
    md_path.write_text(
        (
            "# Customs Code\n\n"
            "Introductory paragraph.\n\n"
            "## Article 1: Scope\n\n"
            "This code governs all customs operations in Uzbekistan.\n\n"
            "## Article 2: Definitions\n\n"
            "Goods means any movable property.\n"
        ),
        encoding="utf-8",
    )
    return md_path


@pytest.mark.asyncio
@pytest.mark.integration
async def test_ingest_markdown_file_writes_chunks(
    ingest_service: LexIngestionService, sample_md: Path
) -> None:
    result = await ingest_service.ingest(sample_md)
    assert result.source_type == "markdown"
    assert result.chunks_written > 0
    assert result.raw_markdown_chars > 0


@pytest.mark.asyncio
@pytest.mark.integration
async def test_ingest_markdown_text_directly(
    ingest_service: LexIngestionService,
) -> None:
    md = "# Title\n\nSome body text here.\n\n## Article 1\n\nMore text."
    result = await ingest_service.ingest_markdown_text(md, source="<test>")
    assert result.chunks_written >= 1


@pytest.mark.asyncio
@pytest.mark.integration
async def test_ingested_chunks_are_searchable(
    ingest_service: LexIngestionService, vector_store, sample_md: Path
) -> None:
    """After ingestion, the content must be retrievable via search_lex."""
    await ingest_service.ingest(sample_md)

    hits = await vector_store.search_lex(
        "What does the customs code say about scope?",
        top_k=5,
    )
    assert hits, "expected at least one hit"
    # The most relevant chunk should mention Article 1 or scope-related content.
    top_text = hits[0].text.lower()
    assert "scope" in top_text or "article 1" in top_text or "customs" in top_text

    # Metadata should preserve the heading hierarchy.
    top_meta = hits[0].metadata
    assert "breadcrumb" in top_meta
    assert "source" in top_meta


@pytest.mark.asyncio
@pytest.mark.integration
async def test_ingest_docx(
    ingest_service: LexIngestionService, tmp_path: Path
) -> None:
    """End-to-end DOCX ingestion."""
    from docx import Document

    docx_path = tmp_path / "law.docx"
    doc = Document()
    doc.add_heading("Customs Code", level=1)
    doc.add_paragraph("Introductory paragraph about customs in Uzbekistan.")
    doc.add_heading("Article 1: Medical Devices", level=2)
    doc.add_paragraph(
        "Medical diagnostic devices with HS code 3822 are duty-free under Resolution 408."
    )
    doc.save(str(docx_path))

    result = await ingest_service.ingest(docx_path)
    assert result.source_type == "docx"
    assert result.chunks_written >= 1
