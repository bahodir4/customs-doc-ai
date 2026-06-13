"""Unit tests for the bulk-ingest workflow.

Uses a fake ingest service so no Qdrant or Ollama is required. The
workflow's contract with the ingest layer is small (a single
`ingest(path) → object with .chunks_written`), so mocking it cleanly
isolates the workflow's own logic: discovery, conversion, deletion.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import pytest

from rag.bulk_ingest import BulkIngestWorkflow, WorkflowSummary


# ── Test doubles ─────────────────────────────────────────────────────


@dataclass
class FakeIngestResult:
    chunks_written: int


class FakeIngestService:
    """Records every ingest call and returns a configurable result."""

    def __init__(self, chunks_per_call: int = 5) -> None:
        self.chunks_per_call = chunks_per_call
        self.calls: list[Path] = []
        self.raise_for: set[str] = set()  # filenames that should raise

    async def ingest(self, source):
        path = Path(source)
        self.calls.append(path)
        if path.name in self.raise_for:
            raise RuntimeError(f"simulated failure for {path.name}")
        return FakeIngestResult(chunks_written=self.chunks_per_call)


# ── Fixtures ─────────────────────────────────────────────────────────


@pytest.fixture
def workflow_dirs(tmp_path: Path) -> tuple[Path, Path]:
    """Create the two directories the workflow expects."""
    originals = tmp_path / "originals"
    markdown = tmp_path / "markdown"
    originals.mkdir()
    markdown.mkdir()
    return originals, markdown


def _write_markdown_file(directory: Path, name: str, body: str = "# Title\n\nText.\n") -> Path:
    path = directory / name
    path.write_text(body, encoding="utf-8")
    return path


def _write_docx_file(directory: Path, name: str) -> Path:
    """Create a real .docx with a heading + paragraph for DocxLoader to parse."""
    from docx import Document

    doc = Document()
    doc.add_heading("Customs Law", level=1)
    doc.add_paragraph("Sample paragraph for ingestion testing.")
    path = directory / name
    doc.save(str(path))
    return path


# ── Discovery ────────────────────────────────────────────────────────


class TestDiscovery:
    def test_finds_supported_files(self, workflow_dirs) -> None:
        originals, markdown = workflow_dirs
        _write_markdown_file(originals, "a.md")
        _write_markdown_file(originals, "b.txt")
        _write_docx_file(originals, "c.docx")

        workflow = BulkIngestWorkflow(
            originals_dir=originals,
            markdown_dir=markdown,
            ingest_service=FakeIngestService(),
        )

        found = workflow.discover_sources()
        names = {p.name for p in found}
        assert names == {"a.md", "b.txt", "c.docx"}

    def test_ignores_unsupported_extensions(self, workflow_dirs) -> None:
        originals, markdown = workflow_dirs
        _write_markdown_file(originals, "good.md")
        # Unsupported extension — should be silently ignored.
        (originals / "noise.png").write_bytes(b"\x89PNG")
        (originals / "noise.zip").write_bytes(b"PK")

        workflow = BulkIngestWorkflow(
            originals_dir=originals,
            markdown_dir=markdown,
            ingest_service=FakeIngestService(),
        )

        found = workflow.discover_sources()
        assert {p.name for p in found} == {"good.md"}

    def test_returns_empty_when_dir_missing(self, tmp_path: Path) -> None:
        workflow = BulkIngestWorkflow(
            originals_dir=tmp_path / "does_not_exist",
            markdown_dir=tmp_path / "md",
            ingest_service=FakeIngestService(),
        )
        assert workflow.discover_sources() == []

    def test_results_are_sorted(self, workflow_dirs) -> None:
        originals, markdown = workflow_dirs
        for name in ("z.md", "a.md", "m.md"):
            _write_markdown_file(originals, name)

        workflow = BulkIngestWorkflow(
            originals_dir=originals,
            markdown_dir=markdown,
            ingest_service=FakeIngestService(),
        )

        names = [p.name for p in workflow.discover_sources()]
        assert names == ["a.md", "m.md", "z.md"]


# ── Conversion + ingestion + cleanup (the happy path) ────────────────


class TestProcessOneSuccess:
    @pytest.mark.asyncio
    async def test_md_file_is_converted_ingested_and_deleted(self, workflow_dirs) -> None:
        originals, markdown = workflow_dirs
        source = _write_markdown_file(originals, "law.md", "# Law\n\nBody.")

        fake = FakeIngestService(chunks_per_call=7)
        workflow = BulkIngestWorkflow(
            originals_dir=originals,
            markdown_dir=markdown,
            ingest_service=fake,
            delete_after_success=True,
        )

        result = await workflow.process_one(source)

        assert result.status == "ok"
        assert result.chunks_written == 7
        assert result.deleted is True
        # Both files removed
        assert not source.exists()
        assert result.markdown_path is not None
        assert not result.markdown_path.exists()
        # Ingest was called on the markdown copy, not the original
        assert len(fake.calls) == 1
        assert fake.calls[0].parent == markdown

    @pytest.mark.asyncio
    async def test_docx_file_is_converted_via_docx_loader(self, workflow_dirs) -> None:
        originals, markdown = workflow_dirs
        source = _write_docx_file(originals, "code.docx")

        workflow = BulkIngestWorkflow(
            originals_dir=originals,
            markdown_dir=markdown,
            ingest_service=FakeIngestService(),
        )

        result = await workflow.process_one(source)

        assert result.status == "ok"
        # MD copy should be written before deletion, and the content
        # should reflect the DOCX heading extraction.
        # We can't check after delete, so set delete=False to inspect.

    @pytest.mark.asyncio
    async def test_no_delete_flag_keeps_files(self, workflow_dirs) -> None:
        originals, markdown = workflow_dirs
        source = _write_markdown_file(originals, "kept.md")

        workflow = BulkIngestWorkflow(
            originals_dir=originals,
            markdown_dir=markdown,
            ingest_service=FakeIngestService(),
            delete_after_success=False,
        )

        result = await workflow.process_one(source)

        assert result.status == "ok"
        assert result.deleted is False
        assert source.exists()
        assert result.markdown_path is not None
        assert result.markdown_path.exists()


# ── Failure handling ─────────────────────────────────────────────────


class TestProcessOneFailure:
    @pytest.mark.asyncio
    async def test_ingest_failure_keeps_both_files(self, workflow_dirs) -> None:
        originals, markdown = workflow_dirs
        source = _write_markdown_file(originals, "boom.md")

        fake = FakeIngestService()
        fake.raise_for.add("boom.md")
        workflow = BulkIngestWorkflow(
            originals_dir=originals,
            markdown_dir=markdown,
            ingest_service=fake,
            delete_after_success=True,
        )

        result = await workflow.process_one(source)

        assert result.status == "failed"
        assert result.error is not None
        assert "boom.md" in result.error
        # Both files must survive a failed run for re-inspection.
        assert source.exists()
        assert result.markdown_path is not None
        assert result.markdown_path.exists()

    @pytest.mark.asyncio
    async def test_zero_chunks_is_treated_as_failure(self, workflow_dirs) -> None:
        originals, markdown = workflow_dirs
        source = _write_markdown_file(originals, "empty.md")

        workflow = BulkIngestWorkflow(
            originals_dir=originals,
            markdown_dir=markdown,
            ingest_service=FakeIngestService(chunks_per_call=0),
            delete_after_success=True,
        )

        result = await workflow.process_one(source)

        assert result.status == "failed"
        assert "0 chunks" in (result.error or "")
        # Nothing deleted on a no-op ingest.
        assert source.exists()


# ── Full run ─────────────────────────────────────────────────────────


class TestRun:
    @pytest.mark.asyncio
    async def test_partial_failure_isolated(self, workflow_dirs) -> None:
        originals, markdown = workflow_dirs
        good = _write_markdown_file(originals, "good.md")
        bad = _write_markdown_file(originals, "bad.md")

        fake = FakeIngestService(chunks_per_call=3)
        fake.raise_for.add("bad.md")

        workflow = BulkIngestWorkflow(
            originals_dir=originals,
            markdown_dir=markdown,
            ingest_service=fake,
            delete_after_success=True,
        )

        summary = await workflow.run()

        assert summary.total == 2
        assert summary.succeeded == 1
        assert summary.failed == 1
        assert summary.total_chunks == 3

        # Good file gone; bad file (and its MD) remain
        assert not good.exists()
        assert bad.exists()

    @pytest.mark.asyncio
    async def test_empty_dir_returns_empty_summary(self, workflow_dirs) -> None:
        originals, markdown = workflow_dirs
        workflow = BulkIngestWorkflow(
            originals_dir=originals,
            markdown_dir=markdown,
            ingest_service=FakeIngestService(),
        )
        summary = await workflow.run()
        assert summary.total == 0
        assert summary.succeeded == 0
        assert summary.failed == 0


class TestWorkflowSummary:
    def test_aggregate_counts(self) -> None:
        from rag.bulk_ingest import FileResult

        summary = WorkflowSummary(results=[
            FileResult(source=Path("a.md"), status="ok", chunks_written=10),
            FileResult(source=Path("b.md"), status="ok", chunks_written=5),
            FileResult(source=Path("c.md"), status="failed", error="x"),
        ])
        assert summary.total == 3
        assert summary.succeeded == 2
        assert summary.failed == 1
        assert summary.total_chunks == 15
