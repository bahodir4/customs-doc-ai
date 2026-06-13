"""Unit tests for source loaders."""
from __future__ import annotations

from pathlib import Path

import pytest

from rag import DocxLoader, MarkdownFileLoader, URLLoader


# ── URLLoader (pure HTML processing — no network) ────────────────────


class TestURLLoaderHtmlConversion:
    def test_strips_script_style_nav(self) -> None:
        loader = URLLoader()
        html = """
        <html>
          <head><style>body{color:red}</style></head>
          <body>
            <nav>Navigation noise</nav>
            <header>Header noise</header>
            <main>
              <h1>Customs Code</h1>
              <p>Real content.</p>
            </main>
            <footer>Footer noise</footer>
            <script>tracking()</script>
          </body>
        </html>
        """
        md = loader._html_to_markdown(html)
        assert "Customs Code" in md
        assert "Real content" in md
        assert "Navigation noise" not in md
        assert "Header noise" not in md
        assert "Footer noise" not in md
        assert "tracking" not in md

    def test_produces_atx_headings(self) -> None:
        loader = URLLoader()
        html = "<html><body><h1>Big</h1><h2>Smaller</h2><p>text</p></body></html>"
        md = loader._html_to_markdown(html)
        assert "# Big" in md
        assert "## Smaller" in md

    def test_collapses_excessive_blank_lines(self) -> None:
        loader = URLLoader()
        text = "para 1\n\n\n\n\n\npara 2"
        cleaned = loader._collapse_whitespace(text)
        # Should have at most 2 blank lines between paragraphs (= 3 newlines).
        assert "\n\n\n\n" not in cleaned


class TestLegalSectionPromotion:
    """The fallback that handles lex.uz pages that style headings via CSS."""

    def test_russian_chapters_promoted_to_h2(self) -> None:
        text = "Глава 1\n\nSome content.\n\nГлава 2\n\nMore content."
        out = URLLoader._promote_legal_sections(text)
        assert "## Глава 1" in out
        assert "## Глава 2" in out

    def test_russian_articles_promoted_to_h3(self) -> None:
        text = "Статья 5\n\nКонтент статьи."
        out = URLLoader._promote_legal_sections(text)
        assert "### Статья 5" in out

    def test_russian_razdel_promoted_to_h1(self) -> None:
        text = "Раздел I\n\nГлава 1\n\nСтатья 1\n\nКонтент."
        out = URLLoader._promote_legal_sections(text)
        assert "# Раздел I" in out
        assert "## Глава 1" in out
        assert "### Статья 1" in out

    def test_uzbek_modda_promoted_to_h3(self) -> None:
        text = "Modda 5\n\nMatn."
        out = URLLoader._promote_legal_sections(text)
        assert "### Modda 5" in out

    def test_uzbek_bob_promoted_to_h2(self) -> None:
        text = "Bob 1\n\nMatn."
        out = URLLoader._promote_legal_sections(text)
        assert "## Bob 1" in out

    def test_english_article_promoted_to_h3(self) -> None:
        text = "Article 5\n\nBody text."
        out = URLLoader._promote_legal_sections(text)
        assert "### Article 5" in out

    def test_existing_headers_are_not_double_promoted(self) -> None:
        text = "## Глава 1\n\nContent."
        out = URLLoader._promote_legal_sections(text)
        # Must not become "## ## Глава 1"
        assert out.count("##") == 1

    def test_non_legal_lines_left_alone(self) -> None:
        text = "Regular paragraph.\n\nAnother one without legal keywords."
        out = URLLoader._promote_legal_sections(text)
        assert "Regular paragraph." in out
        assert "Another one without legal keywords." in out
        assert "#" not in out  # nothing was promoted


# ── MarkdownFileLoader ──────────────────────────────────────────────


@pytest.mark.asyncio
async def test_markdown_file_loader_returns_raw_contents(tmp_path: Path) -> None:
    md_path = tmp_path / "doc.md"
    md_path.write_text("# Header\n\nBody.", encoding="utf-8")

    loader = MarkdownFileLoader()
    content = await loader.load(str(md_path))
    assert content == "# Header\n\nBody."


@pytest.mark.asyncio
async def test_markdown_file_loader_raises_for_missing(tmp_path: Path) -> None:
    loader = MarkdownFileLoader()
    with pytest.raises(FileNotFoundError):
        await loader.load(str(tmp_path / "ghost.md"))


# ── DocxLoader (uses python-docx) ────────────────────────────────────


def _make_docx(path: Path) -> None:
    """Build a tiny .docx file with mixed heading levels and a table."""
    from docx import Document

    doc = Document()
    doc.add_heading("Customs Code", level=1)
    doc.add_paragraph("Introductory paragraph.")
    doc.add_heading("Article 1: Scope", level=2)
    doc.add_paragraph("Body of article 1.")
    doc.add_heading("Part A", level=3)
    doc.add_paragraph("Detail of part A.")
    doc.add_heading("Article 2", level=2)
    doc.add_paragraph("Body of article 2.")

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "HS Code"
    table.rows[0].cells[1].text = "Rate"
    table.rows[1].cells[0].text = "3822"
    table.rows[1].cells[1].text = "0%"

    doc.save(str(path))


@pytest.mark.asyncio
async def test_docx_loader_emits_markdown_headings(tmp_path: Path) -> None:
    docx_path = tmp_path / "law.docx"
    _make_docx(docx_path)

    loader = DocxLoader()
    md = await loader.load(str(docx_path))

    assert "# Customs Code" in md
    assert "## Article 1: Scope" in md
    assert "### Part A" in md
    assert "## Article 2" in md
    # Plain paragraphs survive as plain text
    assert "Body of article 1." in md
    # Table rows are flattened into pipe-separated lines
    assert "HS Code | Rate" in md
    assert "3822 | 0%" in md


@pytest.mark.asyncio
async def test_docx_loader_raises_for_missing(tmp_path: Path) -> None:
    loader = DocxLoader()
    with pytest.raises(FileNotFoundError):
        await loader.load(str(tmp_path / "ghost.docx"))
